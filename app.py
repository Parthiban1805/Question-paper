from flask import Flask, render_template, request, redirect, jsonify, send_file, session
import os
import re
import google.generativeai as genai
from docx import Document
from fpdf import FPDF
from html import escape
import sqlite3
import pandas as py
import threading
import openpyxl
db_lock = threading.Lock()
app = Flask(__name__)
app.secret_key = "logith18801880."

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# Configure Gemini API
# IMPORTANT: The previous API key was blocked by Google due to being leaked.
# Please get a new key from https://aistudio.google.com/app/apikey
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "AIzaSyBrkhMB1Ys_12EctscDW-Bnn0H7IcUQpb4")
genai.configure(api_key=GEMINI_API_KEY)

# Configure generation parameters
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

# The gemini-2.0-flash-exp model has been removed/changed. Switched to gemini-2.5-flash (or gemini-2.0-flash).
model = genai.GenerativeModel(
    model_name="gemini-2.0-flash",
    generation_config=generation_config,
)


def format_questions(text):
    """Parse the text response from Gemini into structured question data."""
    units = {}
    current_unit = None
    current_question = {}
    
    lines = text.strip().split("\n")
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Identify unit
        if re.match(r'^\s*unit:\s*', line, re.IGNORECASE):
            current_unit = line.split(':')[1].strip()
            if current_unit not in units:
                units[current_unit] = []
            continue
        
        # Identify question number
        if re.match(r'^\s*qno:\s*', line, re.IGNORECASE):
            current_question = {'qno': line.split(':')[1].strip()}
            continue
        
        # Identify question text
        if re.match(r'^\s*question:\s*', line, re.IGNORECASE):
            current_question['question'] = line.split(':')[1].strip()
            continue
        
        # Identify marks
        if re.match(r'^\s*marks:\s*', line, re.IGNORECASE):
            current_question['marks'] = line.split(':')[1].strip()
            if 'qno' in current_question and 'question' in current_question:
                # Prevent duplicate questions
                if current_question not in units[current_unit]:
                    units[current_unit].append(current_question)
            current_question = {}  # Reset after storing
    
    return units



@app.route("/upload", methods=["POST"])
def upload_file():
    try:
        file = request.files.get("file")
        subject_name = request.form.get("subject-name")
        course_code = request.form.get("course-code")
        semester = request.form.get("semester")
        exam_time = request.form.get("exam-time")
        total_marks = request.form.get("total-marks")

        if not file or not file.filename.endswith(".docx"):
            return jsonify({"error": "Invalid file format. Please upload a .docx file."}), 400

        file_path = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(file_path)
        
        content = extract_text_from_docx(file_path)
        response_text = generate_questions(content)  # Get clean text response
        
        questions_by_unit = format_questions(response_text)  # Format questions properly
        
        # Store exam details in session
        session['exam_details'] = {
            'subject_name': subject_name,
            'course_code': course_code,
            'semester': semester,
            'exam_time': exam_time,
            'total_marks': total_marks
        }
        
        return render_template(
            "result.html",
            questions_by_unit=questions_by_unit,
            subject_name=subject_name,
            course_code=course_code,
            semester=semester,
            exam_time=exam_time,
            total_marks=total_marks
        )
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def extract_text_from_docx(docx_path):
    """Extracts text from a Word document."""
    text = ""
    document = Document(docx_path)
    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"
    print(text)
    return text

def generate_questions(text_content):
    """Generates questions using Gemini API and ensures valid JSON output."""
    prompt = f"""
    You are a question paper generation AI model.
    You will be given a document about the course content of a college, which contains the topics covered in the specific units.
    You need to generate a question paper for the given document based on the topics covered in the units.
    The questions for each unit should follow this format:

    unit: I 

    qno: 1(a)
    question: What are the four branches of Machine Learning?",
    marks: 2 Marks / CLO1
                    
    qno: 1(b)
    question: Explain the importance of data preprocessing in deep learning. (or) Describe the differences between overfitting and underfitting in Machine Learning.
    marks: 6 Marks / CLO3
                    
    qno: 1(c)
    question: Explain the applications of reinforcement learning in robotics. (or) Discuss the role of activation functions in deep learning.
    marks: 12 Marks / CLO4

    unit: II 

    qno: 2(a)
    question: why do we use pandas?",
    marks: 2 Marks / CLO2
                    
    qno: 2(b)
    question: Explain the importance of python librarries (or) Explain about api integration using gemini.
    marks: 6 Marks / CLO4
                    
    qno: 2(c)
    question: what is the difference between gemini and chatgpt. (or) what is meant by deep seek?.
    marks: 12 Marks / CLO3

    generate this for every unit. most importantly dont give the response in bold text and dont give any extra comment, just what need to given thats all

    ### Course Content:
    {text_content}
    """

    
    try:
        chat_session = model.start_chat(history=[])
        response = chat_session.send_message(prompt)
        print(response.text)
        return response.text
          # Return the raw text response
    except Exception as e:
        print(f"Error generating questions: {str(e)}")
        return str(e)

class QuestionPDF(FPDF):
    first_page = True
    
    def header(self):
        if self.first_page:
            # College Name and Department
            self.set_font('Arial', 'B', 14)
            self.cell(0, 10, 'VET Institute of Arts and Science', 0, 1, 'C')
            
            if hasattr(self, 'exam_details'):
                # Subject and Course Code
                self.set_font('Arial', 'B', 12)
                self.cell(0, 10, f"{self.exam_details['subject_name']} ({self.exam_details['course_code']})", 0, 1, 'C')
                
                # Semester
                self.cell(0, 5, f"Semester: {self.exam_details['semester']}", 0, 1, 'C')
                
                # Time and Marks in the same line
                self.ln(5)
                self.set_font('Arial', '', 10)
                
                # Time on the left
                self.cell(95, 5, f"Time: {self.exam_details['exam_time']} Hours", 0, 0, 'L')
                
                # Maximum marks on the right
                self.cell(95, 5, f"Maximum: {self.exam_details['total_marks']} Marks", 0, 1, 'R')
                
                # Add "Answer all questions" if needed
                self.ln(5)
                self.set_font('Arial', '', 10)
                self.cell(0, 5, "Answer all questions", 0, 1, 'L')
            
            self.ln(10)
            self.first_page = False

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')



@app.route("/save_pdf", methods=["POST"])
def save_pdf():
    try:
        data = request.get_json()
        table_data = data.get('tableData', [])
        
        pdf = QuestionPDF()
        # Set exam details for the PDF
        pdf.exam_details = session.get('exam_details', {})
        
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Define consistent table dimensions
        col_widths = [20, 25, 100, 45]  # Total width = 190
        headers = ['UNIT', 'Q. No.', 'Questions', 'Marks / CLO LEVEL']
        
        def draw_table_borders(y_start, y_end):
            pdf.line(10, y_start, 10, y_end)  # Left border
            pdf.line(30, y_start, 30, y_end)  # After UNIT
            pdf.line(55, y_start, 55, y_end)  # After Q. No.
            pdf.line(155, y_start, 155, y_end)  # After Questions
            pdf.line(200, y_start, 200, y_end)  # Right border
        
        def add_table_header():
            pdf.set_font('Arial', 'B', 10)
            y_start = pdf.get_y()
            
            # Draw header cells
            pdf.set_x(10)
            for i, header in enumerate(headers):
                pdf.cell(col_widths[i], 10, header, 1, 0, 'C')
            pdf.ln()
            
            return y_start
        
        # Process each unit's questions
        for unit_idx, unit_questions in enumerate(table_data):
            if unit_idx > 0:
                if pdf.get_y() + 50 > pdf.page_break_trigger:  # Check if enough space for new table
                    pdf.add_page()
                else:
                    pdf.ln(20)  # Add spacing between tables
            
            y_start = pdf.get_y()
            y_start = add_table_header()
            
            # Process rows
            pdf.set_font('Arial', '', 10)
            for row in unit_questions:
                if pdf.get_y() + 20 > pdf.page_break_trigger:
                    draw_table_borders(y_start, pdf.get_y())
                    pdf.add_page()
                    y_start = pdf.get_y()
                    add_table_header()
                
                # Calculate row height based on question text
                question_text = row[2]
                chars_per_line = 45
                lines_needed = len(question_text) // chars_per_line + 1
                row_height = max(8 * lines_needed, 8)
                
                y_current = pdf.get_y()
                
                # Draw cells with consistent positioning
                pdf.set_xy(10, y_current)
                pdf.cell(col_widths[0], row_height, row[0], 0, 0, 'C')  # UNIT
                pdf.set_xy(30, y_current)
                pdf.cell(col_widths[1], row_height, row[1], 0, 0, 'C')  # Q. No.
                
                # Handle multi-line question text
                pdf.set_xy(55, y_current)
                pdf.multi_cell(col_widths[2], row_height/lines_needed, question_text, 0, 'L')
                
                # Draw marks
                pdf.set_xy(155, y_current)
                pdf.cell(col_widths[3], row_height, row[3], 0, 1, 'C')
                
                # Draw horizontal line
                pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            
            # Draw final table borders for current unit
            draw_table_borders(y_start, pdf.get_y())
        
        # Save PDF
        pdf_path = os.path.join(OUTPUT_FOLDER, "question_paper.pdf")
        pdf.output(pdf_path)
        
        return jsonify({
            "message": "PDF generated successfully!",
            "pdf_url": f"/download/{os.path.basename(pdf_path)}"
        })
        
    except Exception as e:
        return jsonify({"error": str(e)})

def ensure_directory_exists(file_path):
    """Create directory if it doesn't exist"""
    directory = os.path.dirname(file_path)
    if directory and not os.path.exists(directory):
        os.makedirs(directory)

def init_db():
    """Initialize database and create necessary directories"""
    os.makedirs('temp', exist_ok=True)
    
    with db_lock:
        conn = sqlite3.connect('database.db')
        conn.execute("PRAGMA journal_mode=WAL;")
        cursor = conn.cursor()
        cursor.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            full_name TEXT,
            username TEXT,
            email TEXT UNIQUE,
            password TEXT,
            address TEXT,
            phone TEXT
        )
        """)
        conn.commit()
        conn.close()

@app.route('/')
def home():
    return render_template('home.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'GET':
        return render_template('register.html')
    elif request.method == 'POST':
        data = request.json
        full_name = data['full_name']
        username = data['username']
        email = data['email']
        password = data['password']
        address = data['address']
        phone = data['phone']

        try:
            with db_lock:
                conn = sqlite3.connect('database.db')
                cursor = conn.cursor()
                cursor.execute("""
                INSERT INTO users (full_name, username, email, password, address, phone)
                VALUES (?, ?, ?, ?, ?, ?)
                """, (full_name, username, email, password, address, phone))
                conn.commit()
                conn.close()

            excel_path = os.path.join(os.getcwd(), 'users.xlsx')
            save_to_excel(
                [full_name, username, email, password, address, phone],
                excel_path,
                ["Full Name", "Username", "Email", "Password", "Address", "Phone"]
            )

            return jsonify({"message": "Registration successful!"}), 201
        except sqlite3.IntegrityError:
            return jsonify({"error": "Email already exists"}), 400
        except Exception as e:
            return jsonify({"error": str(e)}), 500

def save_to_excel(data, excel_path, headers):
    """Generic function to save data to Excel"""
    try:
        ensure_directory_exists(excel_path)
        if os.path.exists(excel_path):
            wb = openpyxl.load_workbook(excel_path)
            sheet = wb.active
        else:
            wb = openpyxl.Workbook()
            sheet = wb.active
            sheet.append(headers)
        
        sheet.append(data)
        wb.save(excel_path)
        return True
    except Exception as e:
        print(f"Error saving to Excel: {str(e)}")
        return False

@app.route('/logout')
def logout():
    return render_template('home.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'GET':
        return render_template('login.html')
    elif request.method == 'POST':
        data = request.json
        email = data['email']
        password = data['password']

        try:
            conn = sqlite3.connect('database.db')
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM users WHERE email = ? AND password = ?", (email, password))
            user = cursor.fetchone()
            conn.close()

            if user:
                    session['email'] = email  # Store email in session
                    return jsonify({"redirect": '/index'}), 200  # Redirect to `index`
            else:
                    return jsonify({"error": "Invalid email or password"}), 401
        except Exception as e:
                return jsonify({"error": str(e)}), 500

@app.route('/index')
def index():
    if 'email' in session:  # Check if user is logged in
        return render_template('index.html')
    else:
        return redirect('/login')  # Redirect to login if not authenticated

@app.route("/download/<filename>")
def download_file(filename):
    return send_file(
        os.path.join(OUTPUT_FOLDER, filename),
        as_attachment=True,
        download_name=filename
    )

if __name__ == "__main__":
    init_db()
    app.run(debug=True)